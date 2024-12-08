# # utils.py
# import os
# from dotenv import load_dotenv
# from pymongo import MongoClient
# import pdfplumber
# import docx
# from fpdf import FPDF
# import google.generativeai as genai
# from docx import Document
# from docx.shared import Inches

# load_dotenv()

# genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
# model = genai.GenerativeModel("models/gemini-1.5-pro")

# def get_mongo_client():
#     return MongoClient(os.getenv("MONGO_URI"))

# def extract_text_from_file(file):
#     ext = file.name.split(".")[-1].lower()
#     if ext == "pdf":
#         with pdfplumber.open(file) as pdf:
#             return "".join([page.extract_text() for page in pdf.pages])
#     elif ext == "docx":
#         doc = docx.Document(file)
#         return "\n".join([para.text for para in doc.paragraphs])
#     elif ext == "txt":
#         return file.read().decode("utf-8")
#     return None

# def generate_mcqs(text, num_questions, num_options):
#     prompt =  f"""
#     You are an AI assistant helping the user generate multiple-choice questions (MCQs) based on the following text:
#     '{text}'
#     Please generate {num_questions} MCQs from the text. Each question should have: A clear question and,
#     - {num_options} answer options (labeled A, B, C, etc.)
#     - The correct answer clearly indicated
#     Format:
#     ## MCQ
#     [question]
#     A) [option A]
#     B) [option B]
#     ...
#     Correct Answer: [correct option]
#     """

#     response = model.generate_content(prompt).text.strip()
#     return parse_mcqs(response)

# # def parse_mcqs(mcq_text):
# #     mcqs = []
# #     for block in mcq_text.split("## MCQ")[1:]:
# #         lines = block.strip().split("\n")
# #         question = lines[0].replace("Question: ", "")
# #         options = lines[1:-1]
# #         correct_answer = lines[-1].replace("Correct Answer: ", "")
# #         mcqs.append({"question": question, "options": options, "correct_answer": correct_answer})
# #     return mcqs
# def parse_mcqs(mcq_text):
#     mcqs = []
#     for block in mcq_text.split("## MCQ")[1:]:
#         lines = [line.strip() for line in block.strip().split("\n") if line.strip()]  # Clean and filter empty lines
#         if not lines or len(lines) < 3:  # Ensure there's a question and at least two options
#             continue
        
#         question = lines[0].replace("Question: ", "").strip()
#         options = [line.replace(f"{chr(65 + i)})", "").strip() for i, line in enumerate(lines[1:-1])]  # Extract options
#         correct_answer = lines[-1].replace("Correct Answer: ", "").strip()
        
#         mcqs.append({
#             "question": question,
#             "options": options,
#             "correct_answer": correct_answer
#         })
#     return mcqs


# def create_shareable_link(quiz_id):
#     return f"http://localhost:8501/?quiz_id={quiz_id}"

# def generate_pdf(mcqs):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
    
#     for i, mcq in enumerate(mcqs, 1):
#         pdf.cell(200, 10, txt=f"Question {i}: {mcq['question']}", ln=1)
#         for option in mcq['options']:
#             pdf.cell(200, 10, txt=f"- {option}", ln=1)
#         pdf.cell(200, 10, txt=f"Correct Answer: {mcq['correct_answer']}", ln=1)
#         pdf.cell(200, 10, txt="", ln=1)
    
#     return pdf.output(dest='S').encode('latin-1')

# def generate_docx(mcqs):
#     doc = Document()
    
#     for i, mcq in enumerate(mcqs, 1):
#         doc.add_paragraph(f"Question {i}: {mcq['question']}")
#         for option in mcq['options']:
#             doc.add_paragraph(f"- {option}")
#         doc.add_paragraph(f"Correct Answer: {mcq['correct_answer']}")
#         doc.add_paragraph()
    
#     return doc

# def generate_student_result_pdf(student_name, score, total_questions):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)
    
#     pdf.cell(200, 10, txt=f"Student Name: {student_name}", ln=1)
#     pdf.cell(200, 10, txt=f"Score: {score}/{total_questions}", ln=1)
#     # pdf.cell(200, 10, txt=f"Your answer: {answer}", ln=1)
#     # pdf.cell(200, 10, txt=f"Correct answer: {correct_answer}", ln=1)


    
#     return pdf.output(dest='S').encode('latin-1')
# utils.py
import os
from dotenv import load_dotenv
from pymongo import MongoClient
import pdfplumber
import docx
from fpdf import FPDF
import google.generativeai as genai
from docx import Document
from docx.shared import Inches

load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("models/gemini-1.5-pro")

def get_mongo_client():
    return MongoClient(os.getenv("MONGO_URI"))

def extract_text_from_file(file):
    ext = file.name.split(".")[-1].lower()
    if ext == "pdf":
        with pdfplumber.open(file) as pdf:
            return "".join([page.extract_text() for page in pdf.pages])
    elif ext == "docx":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == "txt":
        return file.read().decode("utf-8")
    return None

def generate_mcqs(text, num_questions, num_options):
    prompt =  f"""
    You are an AI assistant helping the user generate multiple-choice questions (MCQs) based on the following text:
    '{text}'
    Please generate {num_questions} MCQs from the text. Each question should have: A clear question and,
    - {num_options} answer options (labeled A, B, C, etc.)
    - The correct answer clearly indicated
    Format:
    ## MCQ
    [question]
    A) [option A]
    B) [option B]
    ...
    Correct Answer: [correct option]
    """

    response = model.generate_content(prompt).text.strip()
    return parse_mcqs(response)

def parse_mcqs(mcq_text):
    if isinstance(mcq_text, list):
        return mcq_text
    
    mcqs = []
    for block in mcq_text.split("## MCQ")[1:]:
        lines = [line.strip() for line in block.strip().split("\n") if line.strip()]
        if not lines or len(lines) < 3:
            continue
        
        question = ' '.join(lines[0].split()[1:]) if lines[0].split()[0].isdigit() else lines[0]
        options = []
        correct_answer = ""
        
        for line in lines:
            if line.startswith("Correct Answer:"):
                correct_answer = line.replace("Correct Answer:", "").strip()
            elif line[0].isalpha() and line[1] == ")":
                options.append(line)
            else:
                # if not question:  # This ensures we capture multi-line questions
                #     question += line + " "
                # else:
                question += " "+line
        
        mcqs.append({
            "question": question.strip(),
            "options": options,
            "correct_answer": correct_answer
        })
    return mcqs

def create_shareable_link(quiz_id):
    base_url = os.environ.get('BASE_URL', 'https://your-app-name.onrender.com')
    return f"{base_url}/?quiz_id={quiz_id}

def generate_pdf(mcqs):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for i, mcq in enumerate(mcqs, 1):
        pdf.cell(200, 10, txt=f"Question {i}: {mcq['question']}", ln=1)
        for option in mcq['options']:
            pdf.cell(200, 10, txt=f"- {option}", ln=1)
        pdf.cell(200, 10, txt=f"Correct Answer: {mcq['correct_answer']}", ln=1)
        pdf.cell(200, 10, txt="", ln=1)
    
    return pdf.output(dest='S').encode('latin-1')

def generate_docx(mcqs):
    doc = Document()
    
    for i, mcq in enumerate(mcqs, 1):
        doc.add_paragraph(f"Question {i}: {mcq['question']}")
        for option in mcq['options']:
            doc.add_paragraph(f"- {option}")
        doc.add_paragraph(f"Correct Answer: {mcq['correct_answer']}")
        doc.add_paragraph()
    
    return doc

def generate_student_result_pdf(student_name, score, total_questions):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    pdf.cell(200, 10, txt=f"Student Name: {student_name}", ln=1)
    pdf.cell(200, 10, txt=f"Score: {score}/{total_questions}", ln=1)
    
    return pdf.output(dest='S').encode('latin-1')
